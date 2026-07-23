"""Convert a simple markdown draft to a Packt-styled .docx deliverable.

Usage: uv run tools/md_to_docx.py <draft.md> <out.docx>

Handles #/##/### headings, - bullets, numbered lists, **bold**, *italic*,
`inline code`, ``` fenced code blocks, > blockquotes (Note/Tip boxes),
and strips HTML comments (draft-only notes stay in the markdown).

Output builds on the template-attached base doc from the author bundle,
so every paragraph carries a named [PACKT] style: HS - ChapterTitle,
HS - Heading 1/2, P0 - Paragraph, L1 - Bullet/Numbered, C0 - CodeBlock,
CS - InlineCode, and I0 - InfoBox / T0 - TipBox.
"""

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches

# the template-attached base doc from the Packt author bundle, vendored here so
# the converter does not depend on the gitignored .tmp layout
PACKT_BASE = Path(__file__).resolve().parent / "templates/packt_base.docx"

SRC, DST = sys.argv[1], sys.argv[2]

STYLES = {
    "title": "HS - ChapterTitle",
    "h1": "HS - Heading 1",
    "h2": "HS - Heading 2",
    "body": "P0 - Paragraph",
    "bullet": "L1 - Bullet",
    "numbered": "L1 - Numbered",
    "code": "C0 - CodeBlock",
    "note": "I0 - InfoBox",
    "tip": "T0 - TipBox",
    "caption": "F0 - FigureCaption",
}
INLINE_CODE_STYLE = "CS - InlineCode"

EMPH_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")


def make_packt_doc():
    """Make a new docx based on the Packt template, clearing the sample content."""
    doc = Document(str(PACKT_BASE))
    # clear the sample content, keep the styles and section setup
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    sect = deepcopy(sect) if sect is not None else None
    for child in list(body):
        body.remove(child)
    if sect is not None:
        body.append(sect)
    return doc


doc = make_packt_doc()


def add_runs(par, text):
    """Add runs to a paragraph for inline formatting: bold, italic, inline code."""
    pos = 0
    for m in EMPH_RE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos : m.start()])
        if m.group(1) is not None:
            par.add_run(m.group(1)).bold = True
        elif m.group(2) is not None:
            par.add_run(m.group(2)).italic = True
        else:
            par.add_run(m.group(3)).style = INLINE_CODE_STYLE
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def add_para(text, kind):
    """Add a paragraph with the given style, and add runs for inline formatting."""
    par = doc.add_paragraph("", style=STYLES[kind])
    add_runs(par, text)
    return par


def add_code_block(block_lines):
    """Add one paragraph per block, manual line breaks preserve spacing."""
    par = doc.add_paragraph("", style=STYLES["code"])
    for i, code_line in enumerate(block_lines):
        if i:
            par.add_run().add_break(WD_BREAK.LINE)
        par.add_run(code_line)


with open(SRC) as f:
    lines = f.read().splitlines()

in_comment = False
in_code = False
code_lines = []
for line in lines:
    stripped = line.strip()
    if in_code:
        if stripped.startswith("```"):
            add_code_block(code_lines)
            code_lines = []
            in_code = False
        else:
            code_lines.append(line.rstrip())
        continue
    if "<!--" in stripped:
        in_comment = "-->" not in stripped
        continue
    if in_comment:
        in_comment = "-->" not in stripped
        continue
    if not stripped:
        continue
    if stripped.startswith("```"):
        in_code = True
        continue
    if m := re.match(r"^!\[(.*)\]\((.+)\)$", stripped):
        caption, img = m.group(1), m.group(2)
        img_path = (Path(SRC).resolve().parent / img).resolve()
        par = doc.add_paragraph("", style="F0 - Figure")
        par.add_run().add_picture(str(img_path), width=Inches(5.8))
        if caption:
            add_para(caption, "caption")
    elif stripped.startswith("# "):
        text = stripped[2:].replace(" — DRAFT for review", "")
        # front matter (Preface) uses MainHeading per the Packt preface template
        if text.strip() == "Preface":
            par = doc.add_paragraph("", style="HS - MainHeading")
            add_runs(par, text)
        else:
            add_para(text, "title")
    elif stripped.startswith("## "):
        add_para(stripped[3:], "h1")
    elif stripped.startswith("### "):
        add_para(stripped[4:], "h2")
    elif stripped.startswith("> "):
        text = stripped[2:]
        kind = "tip" if text.lstrip("*").lower().startswith("tip") else "note"
        add_para(text, kind)
    elif stripped.startswith("- "):
        add_para(stripped[2:], "bullet")
    elif re.match(r"^\d+\. ", stripped):
        add_para(stripped.split(". ", 1)[1], "numbered")
    else:
        add_para(stripped, "body")

if in_code and code_lines:
    add_code_block(code_lines)

doc.save(DST)
print(f"Saved {DST} [Packt styles]")
